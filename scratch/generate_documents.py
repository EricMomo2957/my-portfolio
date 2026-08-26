import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT

def generate_straight_word_doc():
    doc = Document()
    
    # Set equal 1-inch margins on all sides so left margin is a straight line
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_date = p_date.add_run("Aug, 2026\n\n")
    r_date.font.name = 'Times New Roman'
    r_date.font.size = Pt(11)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.line_spacing = 1.15
    p_info.paragraph_format.space_after = Pt(12)
    
    r1 = p_info.add_run("ERIC DOMINIC MOMO\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(11)
    
    r2 = p_info.add_run("Cebu City, Philippines 6000\nmomoe2957@gmail.com\n+63 935 273 7624\n\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    
    r3 = p_info.add_run("The Hiring Manager\n\n")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(11)
    
    r4 = p_info.add_run("Dear Sir/Ma’am,")
    r4.font.name = 'Times New Roman'
    r4.font.size = Pt(11)
    
    paragraphs = [
        "I am writing to express my interest in joining your team. I am a fresh graduate with a Bachelor of Science in Information Technology (BSIT) degree from the University of Cebu – Main Campus where I focused on building practical, real-world applications. I am highly motivated to jumpstart my professional career and am open to any entry-level IT positions—such as Associate Web Developer, Help-Desk, WordPress Assistant or etc.—where I can apply my skills, learn your specific workflows, and grow with your company.",
        "I have gained real-world hands-on experience building systems for actual clients. As part of a 4-member team called KATD SOLUTIONS (4-Member Freelance Team), We developed a Loan Monitoring and Financial Management System for the UC METC Campus. Working alongside another backend developer, we built the API from using Node.js, Express, and PostgreSQL, coding the real-world mathematical formulas cooperatives use to calculate monthly amortizations and track investments.",
        "Additionally, during my internship, I managed daily office paperwork and digital documentation, used WordPress and Elementor to edit client pages, and managed graphic tasks in Canva. I also independently built MentorLog—a system driven by a local database that the company actively used for the daily login and attendance tracking of OJT students. Recently, I also served as the Lead Developer for our Capstone project, ChronoNav, where I managed our code using Git/GitHub, designed layouts in Figma, and helped write our project documentation for the defense.",
        "I am a responsible, cooperative, and coachable person who is ready to learn your workflows. While I am aiming for a Web Developer role, I am completely open to other positions if my skills match a different opening.",
        "I have attached my resume for your review and look forward to the chance to discuss how I can help your team. Thank you for your time and consideration!"
    ]
    
    for text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Fully justify paragraph text so right and left margins are straight!
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        
    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_close.paragraph_format.line_spacing = 1.15
    r_close = p_close.add_run("Sincerely,\n\n\nEric Dominic Momo\nApplicant")
    r_close.font.name = 'Times New Roman'
    r_close.font.size = Pt(11)
    
    os.makedirs("public", exist_ok=True)
    doc.save("public/Eric_Dominic_Momo_Cover_Letter.docx")
    print("Saved public/Eric_Dominic_Momo_Cover_Letter.docx with justified straight margins")

def generate_straight_pdf_doc():
    pdf_filename = "public/Eric_Dominic_Momo_Cover_Letter.pdf"
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # Left-aligned style for headers, dates, and sign-offs
    left_style = ParagraphStyle(
        'DocLeft',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        leading=14.5,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000'),
        spaceAfter=6
    )
    
    # Justified style for body paragraphs so right margin lines up perfectly straight!
    justify_style = ParagraphStyle(
        'DocJustify',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        leading=14.5,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#000000'),
        spaceAfter=10
    )
    
    story = []
    story.append(Paragraph("Aug, 2026", left_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("ERIC DOMINIC MOMO", left_style))
    story.append(Paragraph("Cebu City, Philippines 6000", left_style))
    story.append(Paragraph("momoe2957@gmail.com", left_style))
    story.append(Paragraph("+63 935 273 7624", left_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("The Hiring Manager", left_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Dear Sir/Ma’am,", left_style))
    story.append(Spacer(1, 6))
    
    paragraphs = [
        "I am writing to express my interest in joining your team. I am a fresh graduate with a Bachelor of Science in Information Technology (BSIT) degree from the University of Cebu – Main Campus where I focused on building practical, real-world applications. I am highly motivated to jumpstart my professional career and am open to any entry-level IT positions—such as Associate Web Developer, Help-Desk, WordPress Assistant or etc.—where I can apply my skills, learn your specific workflows, and grow with your company.",
        "I have gained real-world hands-on experience building systems for actual clients. As part of a 4-member team called KATD SOLUTIONS (4-Member Freelance Team), We developed a Loan Monitoring and Financial Management System for the UC METC Campus. Working alongside another backend developer, we built the API from using Node.js, Express, and PostgreSQL, coding the real-world mathematical formulas cooperatives use to calculate monthly amortizations and track investments.",
        "Additionally, during my internship, I managed daily office paperwork and digital documentation, used WordPress and Elementor to edit client pages, and managed graphic tasks in Canva. I also independently built <i>MentorLog</i>—a system driven by a local database that the company actively used for the daily login and attendance tracking of OJT students. Recently, I also served as the Lead Developer for our Capstone project, <i>ChronoNav</i>, where I managed our code using Git/GitHub, designed layouts in Figma, and helped write our project documentation for the defense.",
        "I am a responsible, cooperative, and coachable person who is ready to learn your workflows. While I am aiming for a Web Developer role, I am completely open to other positions if my skills match a different opening.",
        "I have attached my resume for your review and look forward to the chance to discuss how I can help your team. Thank you for your time and consideration!"
    ]
    
    for bp in paragraphs:
        story.append(Paragraph(bp, justify_style))
        
    story.append(Spacer(1, 10))
    story.append(Paragraph("Sincerely,", left_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph("Eric Dominic Momo", left_style))
    story.append(Paragraph("Applicant", left_style))
    
    doc.build(story)
    print("Saved public/Eric_Dominic_Momo_Cover_Letter.pdf with justified straight margins")

if __name__ == "__main__":
    generate_straight_word_doc()
    generate_straight_pdf_doc()
