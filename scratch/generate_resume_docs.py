import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT

def generate_resume_word():
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_name = p_title.add_run("ERIC DOMINIC MOMO\n")
    r_name.bold = True
    r_name.font.name = 'Times New Roman'
    r_name.font.size = Pt(16)
    
    r_contact = p_title.add_run("+63 935 273 7624 | momoe2957@gmail.com | Cebu City, Philippines\n")
    r_contact.font.name = 'Times New Roman'
    r_contact.font.size = Pt(10.5)
    r_contact.bold = True
    
    # Section: CAREER OBJECTIVE
    p_co = doc.add_paragraph()
    p_co.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_co_head = p_co.add_run("CAREER OBJECTIVE ")
    r_co_head.bold = True
    r_co_head.font.name = 'Times New Roman'
    r_co_head.font.size = Pt(11)
    
    r_co_body = p_co.add_run("A dedicated, adaptable and collaborative IT fresh graduate eager to start a career as Design Associate. I gained practical experience building functional web applications and delivering solutions for real-world clients. I am looking to bring my technical skills and proactive work ethic to your team, where I can contribute to projects while continuing to grow as a professional.\n")
    r_co_body.font.name = 'Times New Roman'
    r_co_body.font.size = Pt(10.5)
    
    # Section: TECHNICAL SKILLS
    p_ts = doc.add_paragraph()
    r_ts_head = p_ts.add_run("TECHNICAL SKILLS")
    r_ts_head.bold = True
    r_ts_head.font.name = 'Times New Roman'
    r_ts_head.font.size = Pt(12)
    
    skills = [
        ("Languages: ", "HTML, CSS, TypeScript, PHP, Python, Java, MySQL, PostgresSQL"),
        ("Frameworks & Libraries: ", "Node.js, Express.js, React, Bootstrap, Vue.js"),
        ("Tools & Platforms: ", "Git, GitHub, MySQL, XAMPP, VS Code, Figma, Canva, PhotoShop, Photopea, WordPress, Elementor, Microsoft Office (Excel, PowerPoint)"),
        ("AI & Automation: ", "Flowise (Logic-based AI Chatbot Development), Professional IT- Elective in Artificial Intelligence"),
        ("Methodologies: ", "OOP, SOLID & SEO Principles, MVC Architecture")
    ]
    for label, val in skills:
        p_s = doc.add_paragraph()
        p_s.paragraph_format.space_after = Pt(2)
        r_l = p_s.add_run("• " + label)
        r_l.bold = True
        r_l.font.name = 'Times New Roman'
        r_l.font.size = Pt(10.5)
        r_v = p_s.add_run(val)
        r_v.font.name = 'Times New Roman'
        r_v.font.size = Pt(10.5)
        
    # Section: PROFESSIONAL PROJECTS & EXPERIENCE
    p_exp_h = doc.add_paragraph()
    p_exp_h.paragraph_format.space_before = Pt(8)
    r_exp = p_exp_h.add_run("PROFESSIONAL PROJECTS & EXPERIENCE")
    r_exp.bold = True
    r_exp.font.name = 'Times New Roman'
    r_exp.font.size = Pt(12)
    
    # Project 1
    p_p1 = doc.add_paragraph()
    r_p1_t = p_p1.add_run("Loan Monitoring and Financial Management System | May 2025 – Present\n")
    r_p1_t.bold = True
    r_p1_t.font.name = 'Times New Roman'
    r_p1_t.font.size = Pt(11)
    
    r_p1_c = p_p1.add_run("KATD SOLUTIONS (4-Member Freelance Team) | Client: UC METC Campus, Cebu City")
    r_p1_c.italic = True
    r_p1_c.font.name = 'Times New Roman'
    r_p1_c.font.size = Pt(10.5)
    
    bullets1 = [
        "Worked with another developer to build the backend API from scratch using Node.js and Express to handle math formulas for monthly payments, making loan processing and tracking much faster.",
        "Focused on using automation to improve daily operations, making it easy for the client to get accurate, up-to-date financial information and transaction reports whenever they need them.",
        "Co-designed and built the database using PostgreSQL to centralize member records, move loan processes online, and cut down on manual paperwork and errors."
    ]
    for b in bullets1:
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pb.paragraph_format.space_after = Pt(3)
        rb = pb.add_run("• " + b)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10.5)
        
    # Project 2
    p_p2 = doc.add_paragraph()
    r_p2_t = p_p2.add_run("On-the-Job Trainee (IT Admin Assistant & Web Development) | Feb 2026 – April 2026\n")
    r_p2_t.bold = True
    r_p2_t.font.name = 'Times New Roman'
    r_p2_t.font.size = Pt(11)
    
    r_p2_c = p_p2.add_run("CoreLogic Consulting & System, Inc. | Gov. M. Cuenco Avenue, Cebu City")
    r_p2_c.italic = True
    r_p2_c.font.name = 'Times New Roman'
    r_p2_c.font.size = Pt(10.5)
    
    bullets2 = [
        "Handled daily office paperwork and digital documentation and used WordPress with Elementor to design, edit, and update different web pages for real clients.",
        "Used Canva and photoshop to create graphic assets, including video editing, logos, and pictures, while independently building Mentor Log—a local-base system the company actively uses to automate trainee logs and mentorship tracking.",
        "Managed local data safety by regularly backing up client files and organizing company digital records to prevent data loss."
    ]
    for b in bullets2:
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pb.paragraph_format.space_after = Pt(3)
        rb = pb.add_run("• " + b)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10.5)

    # Project 3
    p_p3 = doc.add_paragraph()
    r_p3_t = p_p3.add_run("CHRONONAV – Campus Navigation & Scheduling App (Capstone Lead Developer) | Aug 2025 - Dec 2025")
    r_p3_t.bold = True
    r_p3_t.font.name = 'Times New Roman'
    r_p3_t.font.size = Pt(11)
    
    bullets3 = [
        "Served as the Lead Developer, managing the code using Git/GitHub and ensuring a stable backend to connect campus and navigation data.",
        "Integrated OCR technology to turn scanned student study loads into digital schedules, while helping write the project documentation and presenting during the defense."
    ]
    for b in bullets3:
        pb = doc.add_paragraph()
        pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pb.paragraph_format.space_after = Pt(3)
        rb = pb.add_run("• " + b)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10.5)

    # Section: EDUCATION
    p_edu_h = doc.add_paragraph()
    p_edu_h.paragraph_format.space_before = Pt(8)
    r_edu = p_edu_h.add_run("EDUCATION")
    r_edu.bold = True
    r_edu.font.name = 'Times New Roman'
    r_edu.font.size = Pt(12)
    
    p_edu_c = doc.add_paragraph()
    r_ec = p_edu_c.add_run("University of Cebu – Main Campus | Cebu City, 6000 Philippines")
    r_ec.bold = True
    r_ec.font.name = 'Times New Roman'
    r_ec.font.size = Pt(11)
    
    p_eb = doc.add_paragraph()
    r_eb = p_eb.add_run("• Bachelor of Science in Information Technology | June 2026")
    r_eb.font.name = 'Times New Roman'
    r_eb.font.size = Pt(10.5)

    # Section: CERTIFICATIONS
    p_cert_h = doc.add_paragraph()
    p_cert_h.paragraph_format.space_before = Pt(8)
    r_cert = p_cert_h.add_run("CERTIFICATIONS")
    r_cert.bold = True
    r_cert.font.name = 'Times New Roman'
    r_cert.font.size = Pt(12)
    
    p_c1 = doc.add_paragraph()
    r_cc = p_c1.add_run("Cisco Certified Network Associate (CCNAv7)")
    r_cc.bold = True
    r_cc.font.name = 'Times New Roman'
    r_cc.font.size = Pt(11)
    
    cbull = [
        "Switching, Routing, and Wireless Essentials | Jan 2025",
        "Introduction to Networks | June 2024"
    ]
    for b in cbull:
        pb = doc.add_paragraph()
        rb = pb.add_run("• " + b)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10.5)

    os.makedirs("public", exist_ok=True)
    doc.save("public/Eric_Dominic_Momo_Resume.docx")
    print("Saved public/Eric_Dominic_Momo_Resume.docx")

def generate_resume_pdf():
    pdf_filename = "public/Eric_Dominic_Momo_Resume.pdf"
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'ResTitle',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=16,
        leading=20,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000')
    )
    
    contact_style = ParagraphStyle(
        'ResContact',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=10.5,
        leading=14,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000'),
        spaceAfter=10
    )

    heading_style = ParagraphStyle(
        'ResHeading',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=12,
        leading=16,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000'),
        spaceBefore=10,
        spaceAfter=4
    )
    
    body_justify = ParagraphStyle(
        'ResBodyJustify',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10.5,
        leading=14,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor('#000000'),
        spaceAfter=3
    )

    body_left = ParagraphStyle(
        'ResBodyLeft',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10.5,
        leading=14,
        alignment=TA_LEFT,
        textColor=colors.HexColor('#000000'),
        spaceAfter=3
    )

    story = []
    story.append(Paragraph("ERIC DOMINIC MOMO", title_style))
    story.append(Paragraph("+63 935 273 7624 | momoe2957@gmail.com | Cebu City, Philippines", contact_style))
    
    # Career Objective
    co_text = "<b>CAREER OBJECTIVE</b> A dedicated, adaptable and collaborative IT fresh graduate eager to start a career as Design Associate. I gained practical experience building functional web applications and delivering solutions for real-world clients. I am looking to bring my technical skills and proactive work ethic to your team, where I can contribute to projects while continuing to grow as a professional."
    story.append(Paragraph(co_text, body_justify))
    story.append(Spacer(1, 6))

    # Technical Skills
    story.append(Paragraph("TECHNICAL SKILLS", heading_style))
    skills = [
        "<b>• Languages:</b> HTML, CSS, TypeScript, PHP, Python, Java, MySQL, PostgresSQL",
        "<b>• Frameworks & Libraries:</b> Node.js, Express.js, React, Bootstrap, Vue.js",
        "<b>• Tools & Platforms:</b> Git, GitHub, MySQL, XAMPP, VS Code, Figma, Canva, PhotoShop, Photopea, WordPress, Elementor, Microsoft Office (Excel, PowerPoint)",
        "<b>• AI & Automation:</b> Flowise (Logic-based AI Chatbot Development), Professional IT- Elective in Artificial Intelligence",
        "<b>• Methodologies:</b> OOP, SOLID & SEO Principles, MVC Architecture"
    ]
    for s in skills:
        story.append(Paragraph(s, body_left))
    story.append(Spacer(1, 6))

    # Professional Projects & Experience
    story.append(Paragraph("PROFESSIONAL PROJECTS & EXPERIENCE", heading_style))
    
    story.append(Paragraph("<b>Loan Monitoring and Financial Management System</b> | May 2025 – Present", body_left))
    story.append(Paragraph("<i>KATD SOLUTIONS (4-Member Freelance Team) | Client: UC METC Campus, Cebu City</i>", body_left))
    bullets1 = [
        "• Worked with another developer to build the backend API from scratch using Node.js and Express to handle math formulas for monthly payments, making loan processing and tracking much faster.",
        "• Focused on using automation to improve daily operations, making it easy for the client to get accurate, up-to-date financial information and transaction reports whenever they need them.",
        "• Co-designed and built the database using PostgreSQL to centralize member records, move loan processes online, and cut down on manual paperwork and errors."
    ]
    for b in bullets1:
        story.append(Paragraph(b, body_justify))
    story.append(Spacer(1, 4))

    story.append(Paragraph("<b>On-the-Job Trainee (IT Admin Assistant & Web Development)</b> | Feb 2026 – April 2026", body_left))
    story.append(Paragraph("<i>CoreLogic Consulting & System, Inc. | Gov. M. Cuenco Avenue, Cebu City</i>", body_left))
    bullets2 = [
        "• Handled daily office paperwork and digital documentation and used WordPress with Elementor to design, edit, and update different web pages for real clients.",
        "• Used Canva and photoshop to create graphic assets, including video editing, logos, and pictures, while independently building Mentor Log—a local-base system the company actively uses to automate trainee logs and mentorship tracking.",
        "• Managed local data safety by regularly backing up client files and organizing company digital records to prevent data loss."
    ]
    for b in bullets2:
        story.append(Paragraph(b, body_justify))
    story.append(Spacer(1, 4))

    story.append(Paragraph("<b>CHRONONAV – Campus Navigation & Scheduling App (Capstone Lead Developer)</b> | Aug 2025 - Dec 2025", body_left))
    bullets3 = [
        "• Served as the Lead Developer, managing the code using Git/GitHub and ensuring a stable backend to connect campus and navigation data.",
        "• Integrated OCR technology to turn scanned student study loads into digital schedules, while helping write the project documentation and presenting during the defense."
    ]
    for b in bullets3:
        story.append(Paragraph(b, body_justify))
    story.append(Spacer(1, 6))

    # Education
    story.append(Paragraph("EDUCATION", heading_style))
    story.append(Paragraph("<b>University of Cebu – Main Campus</b> | Cebu City, 6000 Philippines", body_left))
    story.append(Paragraph("• Bachelor of Science in Information Technology | <b>June 2026</b>", body_left))
    story.append(Spacer(1, 6))

    # Certifications
    story.append(Paragraph("CERTIFICATIONS", heading_style))
    story.append(Paragraph("<b>Cisco Certified Network Associate (CCNAv7)</b>", body_left))
    story.append(Paragraph("• Switching, Routing, and Wireless Essentials | Jan 2025", body_left))
    story.append(Paragraph("• Introduction to Networks | June 2024", body_left))

    doc.build(story)
    print("Saved public/Eric_Dominic_Momo_Resume.pdf")

if __name__ == "__main__":
    generate_resume_word()
    generate_resume_pdf()
